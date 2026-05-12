"""Generate trifold_board.docx from the 8 talking-point sections."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page setup: landscape letter ──────────────────────────────────────────────
section = doc.sections[0]
section.orientation = WD_ORIENT.LANDSCAPE
section.page_width  = Inches(11)
section.page_height = Inches(8.5)
section.top_margin    = Inches(0.6)
section.bottom_margin = Inches(0.5)
section.left_margin   = Inches(0.7)
section.right_margin  = Inches(0.7)

# ── Styles helpers ─────────────────────────────────────────────────────────────
def set_run_color(run, hex_str):
    r, g, b = int(hex_str[0:2],16), int(hex_str[2:4],16), int(hex_str[4:6],16)
    run.font.color.rgb = RGBColor(r, g, b)

def heading(doc, number, title, subtitle, color_hex):
    """Numbered page heading."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    num_run = p.add_run(f"{number}  ")
    num_run.bold = True
    num_run.font.size = Pt(28)
    set_run_color(num_run, color_hex)
    title_run = p.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(22)
    set_run_color(title_run, color_hex)
    p.paragraph_format.space_after = Pt(2)

    sub = doc.add_paragraph(subtitle)
    sub.runs[0].italic = True
    sub.runs[0].font.size = Pt(11)
    set_run_color(sub.runs[0], "555555")
    sub.paragraph_format.space_after = Pt(8)

    # Divider
    div = doc.add_paragraph()
    div.paragraph_format.space_after = Pt(6)
    pPr = div._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)

def section_title(doc, text, color_hex):
    p = doc.add_paragraph()
    r = p.add_run(text.upper())
    r.bold = True
    r.font.size = Pt(10)
    set_run_color(r, color_hex)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)

def bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.2)
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(text)
    r.font.size = Pt(10)

def body_text(doc, text):
    p = doc.add_paragraph(text)
    p.runs[0].font.size = Pt(10)
    p.paragraph_format.space_after = Pt(4)

def callout(doc, text, color_hex):
    p = doc.add_paragraph()
    r = p.add_run(f"▶  {text}")
    r.italic = True
    r.font.size = Pt(10)
    set_run_color(r, color_hex)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)

def page_break(doc):
    doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PAGE 1 – What Is LiDAR?
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, "01", "What Is LiDAR?", "Understanding the core sensing technology behind this project", "1a4fa0")

section_title(doc, "Definition", "1a4fa0")
body_text(doc, "LiDAR stands for Light Detection and Ranging. It works by emitting laser pulses and measuring the time it takes for each pulse to bounce back from a surface. That travel time is converted into a precise distance measurement — the same principle as sonar, but using light instead of sound.")

section_title(doc, "How a Distance Reading Becomes a 3D Point", "1a4fa0")
for item in [
    "The sensor fires a laser and measures the round-trip time to the target.",
    "Distance (d) is calculated:  d = (c × t) / 2",
    "An orientation reading (from the IMU) provides a direction vector.",
    "Together they place a point in 3D space: (x, y, z)",
    "Thousands of these points build a point cloud — a 3D map.",
]:
    bullet(doc, item)

section_title(doc, "Real-World Uses", "1a4fa0")
for item in ["Self-driving cars", "Robotics & navigation", "Architecture & surveying",
             "Forestry & geology", "AR/VR spatial computing", "Archaeology"]:
    bullet(doc, item)

callout(doc, "This project makes LiDAR accessible at low cost — from a $30 sensor to a live interactive 3D point cloud.", "1a4fa0")

section_title(doc, "Key Stats", "1a4fa0")
for item in ["Sample rate: 250 Hz", "Effective range: 0.2 – 8 m", "Typical accuracy: ±2 cm", "Output: 3D point cloud"]:
    bullet(doc, item)

# ══════════════════════════════════════════════════════════════════════════════
# PAGE 2 – Hardware Components
# ══════════════════════════════════════════════════════════════════════════════
page_break(doc)
heading(doc, "02", "Hardware Components", "Every physical part of the scanner and what it contributes", "0d7377")

components = [
    ("TF-Luna LiDAR",        "Measures distance via time-of-flight laser",                     "UART (serial)",   "250 Hz, 0.2–8 m; returns dist, signal strength & temp"),
    ("BNO055 IMU",           "Provides 9-DOF orientation as a quaternion for point direction", "I²C (bus 1)",     "On-board sensor fusion; outputs qw, qx, qy, qz directly"),
    ("Raspberry Pi 5",       "Central compute — reads sensors, processes data, sends UDP",     "GPIO / USB",      "Runs Python tools, hosts the HAL, streams to visualization"),
    ("Pi Camera Rev 1.3",    "Detects ArUco visual markers for precise scanner positioning",   "CSI ribbon",      "Used with OpenCV for pose estimation; replaced Pi Camera 3"),
    ("3D Printed Housing",   "Holds all electronics in a stable handheld form factor",         "—",               "Designed in Rhino 3D; maintains sensor alignment during scans"),
    ("ArUco Markers",        "Visual reference points to anchor the scanner in space",          "Camera",          "Reduce IMU drift; provide absolute position in the scene"),
]

for name, purpose, interface, detail in components:
    p = doc.add_paragraph()
    r = p.add_run(f"{name}  |  ")
    r.bold = True
    r.font.size = Pt(10)
    set_run_color(r, "0d7377")
    p.add_run(f"{purpose}   [{interface}]   {detail}").font.size = Pt(10)
    p.paragraph_format.space_after = Pt(3)

callout(doc, "Each component was validated with isolated Python test scripts before full integration — catching wiring and config issues early.", "0d7377")

section_title(doc, "Total Component Cost", "0d7377")
body_text(doc, "TF-Luna (~$30), BNO055 breakout (~$10), Raspberry Pi 5 (~$80), Pi Camera (~$25). The system proves that capable spatial scanning does not require expensive commercial hardware.")

# ══════════════════════════════════════════════════════════════════════════════
# PAGE 3 – Physical Construction
# ══════════════════════════════════════════════════════════════════════════════
page_break(doc)
heading(doc, "03", "Physical Construction", "Wiring, CAD design, 3D printing, and final assembly", "1b2a6b")

section_title(doc, "Wiring the Scanner", "1b2a6b")
for item in [
    "TF-Luna → Pi UART: TX→RX (GPIO 14/15), 5 V power & GND",
    "BNO055 → Pi I²C: SDA (GPIO 2), SCL (GPIO 3), 3.3 V & GND",
    "Pi Camera: CSI ribbon connector to Pi 5 camera port",
    "UART and I²C enabled in raspi-config; serial console disabled",
    "User added to dialout group for serial port access",
]:
    bullet(doc, item)

section_title(doc, "CAD & 3D Printing", "1b2a6b")
for item in [
    "Housing modeled in Rhino 3D (.3dm files)",
    "Designed to hold TF-Luna, BNO055, Pi 5, and camera together",
    "Sensor alignment baked into the geometry — no re-calibration needed after assembly",
    "Printed while software development ran in parallel to save time",
]:
    bullet(doc, item)

section_title(doc, "Assembly Steps", "1b2a6b")
for i, item in enumerate([
    "Wire TF-Luna to Pi and validate with tfluna_read.py",
    "Wire BNO055 to Pi and validate with bno055_quat_read.py",
    "Connect Pi Camera and verify image capture",
    "Run camera calibration and capture ArUco marker images",
    "Install components into 3D printed housing",
    "Re-run all sensor tests as an integrated unit",
], 1):
    bullet(doc, f"Step {i}: {item}")

callout(doc, "The housing was designed so the LiDAR sensor, IMU, and camera maintain a fixed spatial relationship — making orientation math between sensors consistent.", "1b2a6b")

# ══════════════════════════════════════════════════════════════════════════════
# PAGE 4 – Software Architecture
# ══════════════════════════════════════════════════════════════════════════════
page_break(doc)
heading(doc, "04", "Software Architecture", "Python tools, HAL interfaces, and the visualization stack", "1c6e3d")

section_title(doc, "Python Side (Raspberry Pi)", "1c6e3d")
for item in [
    "HAL (Hardware Abstraction Layer) — clean interfaces; any sensor can be swapped or mocked without changing other code",
    "tfluna_read.py — UART reader for raw distance frames",
    "bno055_quat_read.py — I²C reader for quaternion orientation",
    "aruco_pose_demo.py — camera-based marker detection and pose estimation",
    "send_sensor_data_udp.py — fuses all sensor streams and broadcasts UDP packets",
]:
    bullet(doc, item)

section_title(doc, "Testing & Mocking", "1c6e3d")
for item in [
    "All HAL interfaces have mock implementations for offline testing",
    "tests/test_hal.py verifies sensor logic without hardware",
    "Mock output logs in test_outputs/ used for data format validation",
]:
    bullet(doc, item)

section_title(doc, "Visualization Side (Unity)", "1c6e3d")
for item in [
    "Unity (C#) — rendering target; validated the full data pipeline",
    "Custom GLSL point cloud shader for real-time rendering",
    "Debug overlay for live diagnostics during scanning",
]:
    bullet(doc, item)

callout(doc, "The HAL design means the scanner's software can be tested on any computer — not just the Pi — which dramatically speeds up development.", "1c6e3d")

section_title(doc, "Data Flow", "1c6e3d")
body_text(doc, "TF-Luna (UART) ──→ Pi Python tools ──→ Coordinate transform ──→ UDP packet ──→ Unity receiver ──→ Point Cloud render\nBNO055 (I²C) ──→ ↗                  ArUco Camera ──→ ↗ (scanner origin fusion)")

# ══════════════════════════════════════════════════════════════════════════════
# PAGE 5 – Initial Testing
# ══════════════════════════════════════════════════════════════════════════════
page_break(doc)
heading(doc, "05", "Initial Testing", "Validating every component before full system integration", "5a2d82")

section_title(doc, "TF-Luna Raw Output Sample", "5a2d82")
body_text(doc, "1770662039.648  dist_cm=0   strength=128  temp_c=39.00\n1770662039.658  dist_cm=0   strength=128  temp_c=39.00\n1770662039.669  dist_cm=42  strength=131  temp_c=39.00")
body_text(doc, "Each frame includes a Unix timestamp, measured distance, return signal strength (quality), and on-chip temperature.")

section_title(doc, "BNO055 Raw Output Sample", "5a2d82")
body_text(doc, "1770662963.567  qw=1.000 qx=0.000 qy=-0.000 qz=0.000\n1770662963.669  qw=0.368 qx=-0.492 qy=0.788 qz=0.000\n1770662963.771  qw=0.368 qx=-0.492 qy=0.788 qz=0.000")
body_text(doc, "Quaternions describe full 3D orientation with no gimbal lock. The BNO055's internal fusion removes the need for manual Euler math.")

section_title(doc, "Incremental Test Strategy", "5a2d82")
for item in [
    "Each sensor tested in isolation before integration",
    "Hardware faults caught early via per-component scripts",
    "UDP stream validated before connecting to Unity",
    "ArUco pipeline tested independently from IMU stream",
    "Camera calibration verified with dedicated image capture tool",
]:
    bullet(doc, item)

callout(doc, "Testing components separately first eliminated wiring and config issues early, so final integration debugging focused only on software logic.", "5a2d82")

# ══════════════════════════════════════════════════════════════════════════════
# PAGE 6 – Sending Data with UDP
# ══════════════════════════════════════════════════════════════════════════════
page_break(doc)
heading(doc, "06", "Sending Data with UDP", "Bridging the Raspberry Pi scanner to the visualization environment in real time", "2e4057")

section_title(doc, "Why UDP?", "2e4057")
for item in [
    "Low overhead — no handshaking or connection state",
    "Low latency — packets arrive as fast as the network allows",
    "Fire-and-forget — acceptable if a few frames are dropped; point clouds are naturally tolerant of gaps",
    "Easy to receive — Unity and Python both have simple UDP socket APIs",
]:
    bullet(doc, item)

section_title(doc, "Packet Contents", "2e4057")
for item in [
    "tfluna — distance (cm), signal strength, temperature",
    "bno055 — quaternion orientation (qw, qx, qy, qz)",
    "aruco — detected marker IDs and pose vectors (when camera is active)",
    "scanner_pose — fused position (from ArUco) + orientation (from IMU)",
    "pos_m — final 3D endpoint in world space, ready to plot",
]:
    bullet(doc, item)

section_title(doc, "Transmission Pipeline", "2e4057")
for item in [
    "Pi reads TF-Luna and BNO055 continuously",
    "ArUco detector sends marker poses to local port 6006",
    "Aggregator fuses all three streams into one JSON packet",
    "Packet broadcast to visualization host on port 5005",
    "Unity receiver parses JSON and queues the new point",
]:
    bullet(doc, item)

callout(doc, "The Pi owns the coordinate math. The visualization engine just plots what arrives — keeping the rendering side simple and engine-agnostic.", "2e4057")

section_title(doc, "Frame Mode Tag", "2e4057")
body_text(doc, "Each packet includes a frame mode tag (unity_frame / raw) so the receiving engine knows whether coordinate conversion has already been applied by the Pi — eliminating double-correction bugs.")

# ══════════════════════════════════════════════════════════════════════════════
# PAGE 7 – Point Cloud Generation
# ══════════════════════════════════════════════════════════════════════════════
page_break(doc)
heading(doc, "07", "Point Cloud Generation", "Turning raw sensor samples into a navigable 3D spatial map", "b53a1e")

section_title(doc, "How Each Point Is Computed", "b53a1e")
for item in [
    "1. Read distance — TF-Luna returns distance d to the nearest surface along the sensor axis.",
    "2. Read orientation — BNO055 returns a unit quaternion q describing where the scanner is pointing.",
    "3. Build direction vector — rotate the scanner's forward vector (0,0,1) by q to get a world-space direction.",
    "4. Compute endpoint — multiply the direction vector by d to get the absolute 3D position of the surface.",
    "5. Accumulate — the new point is appended to the cloud buffer and rendered in real time.",
]:
    bullet(doc, item)

section_title(doc, "Coordinate System Challenges", "b53a1e")
for item in [
    "TF-Luna, BNO055, and Unity each use different axis conventions",
    "Iterative live testing used to determine axis flips, swaps, and yaw offsets",
    "Final calibrated settings: Flip Y, Swap Y↔Z, Invert Z, Yaw correction −180°",
    "Frame mode tagging prevents double-correction between Pi and engine",
]:
    bullet(doc, item)

section_title(doc, "Rendering", "b53a1e")
for item in ["Custom GLSL point cloud shader", "Points colored by depth or height",
             "GPU-side rendering for performance", "Renders in Unity in real time"]:
    bullet(doc, item)

callout(doc, "ArUco markers improve accuracy by anchoring the scanner's 3D position. Without them, small IMU drift causes scans to 'float' over time.", "b53a1e")

# ══════════════════════════════════════════════════════════════════════════════
# PAGE 8 – Conclusion & Lessons Learned
# ══════════════════════════════════════════════════════════════════════════════
page_break(doc)
heading(doc, "08", "Conclusion & Lessons Learned", "What was built, what was gained, and where it goes next", "1a1a2e")

section_title(doc, "What Was Delivered", "1a1a2e")
for item in [
    "Fully wired handheld LiDAR scanner in a custom 3D printed housing",
    "Python sensor stack with HAL interfaces and hardware mock support",
    "ArUco camera pipeline for accurate scanner pose estimation",
    "UDP sensor fusion stream combining LiDAR, IMU, and camera data",
    "Unity point cloud viewer — validated end-to-end pipeline",
]:
    bullet(doc, item)

section_title(doc, "Key Lessons", "1a1a2e")
for item in [
    "Test incrementally — isolate components first, integrate second",
    "Pivot strategically — when a camera driver blocked progress, shifting focus to Unity visualization kept the project moving",
    "Coordinate systems matter — every engine and sensor has its own axes; mapping them correctly is non-trivial",
    "Sensor fusion > single-source — combining IMU + ArUco produces far more stable results than either alone",
]:
    bullet(doc, item)

section_title(doc, "Skills Developed", "1a1a2e")
for item in [
    "Embedded hardware integration (UART, I²C, GPIO)",
    "Computer vision (camera calibration, ArUco detection)",
    "3D math: quaternions, coordinate transforms, point clouds",
    "Network programming (UDP socket design)",
    "Game engine development (Unity C#)",
    "CAD design and 3D printing",
]:
    bullet(doc, item)

callout(doc, "This project combined hardware, software, networking, computer vision, and 3D math into a single working pipeline — demonstrating what a computer science education looks like in practice.", "1a1a2e")

section_title(doc, "Future Directions", "1a1a2e")
for item in [
    "SLAM — simultaneous localization and mapping during movement",
    "Point cloud export and mesh reconstruction",
    "Smoothed sensor fusion with complementary or Kalman filtering",
]:
    bullet(doc, item)

# ── Save ───────────────────────────────────────────────────────────────────────
out = "presentation/trifold_board.docx"
doc.save(out)
print(f"Saved: {out}")
